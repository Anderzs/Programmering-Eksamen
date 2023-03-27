# PYTHON-DOCX
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

#JSON
from json import load, dump

# ENUM
from enum import Enum

#SYSTEM
from sys import platform
import os

#DATETIME
from datetime import datetime

#ARGUMENTPARSER
import argparse


class Fag(Enum):
    FYSIK = {"skabelon": "skabeloner/fysik.json", "output": "output/"}
    KEMI = {"skabelon": "skabeloner/kemi.json", "output": "output/"}
    TEKNOLOGI = {"skabelon": "skabeloner/teknologi.json", "output": "output/"}


class Journal:
    def __init__(self, config: dict) -> None:
        self.title = config['title']
        self.forside = config['front']
        self.save_path = config['out'] if config['out'] else "output\\"
        self.config = config

        self.settings = self.get_content(path="settings.json")
        _fag: str = config['subject']
        match _fag.upper():
            case "FYSIK":
                self.fag = Fag.FYSIK
            case "KEMI":
                self.fag = Fag.KEMI
            case "TEKNOLOGI":
                self.fag = Fag.TEKNOLOGI
            case _:
                raise ValueError("Kunne ikke finde et eksisterende fag (fysik, kemi, teknologi)")
        
        
    def get_content(self, path: str) -> dict:
        with open(path, 'r', encoding='utf-8') as f:
            return load(f)
        
    def is_url(self, url: str) -> bool:
        # Taget fra Django
        # https://github.com/django/django/blob/stable/1.3.x/django/core/validators.py#L45
        import re
        regex = re.compile(
            r'^(?:http|ftp)s?://' # http:// or https://
            r'(?:(?:[A-Z0-9](?:[A-Z0-9-]{0,61}[A-Z0-9])?\.)+(?:[A-Z]{2,6}\.?|[A-Z0-9-]{2,}\.?)|' #domain...
            r'localhost|' #localhost...
            r'\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})' # ...or ip
            r'(?::\d+)?' # optional port
            r'(?:/?|[/?]\S+)$', re.IGNORECASE)

        return re.match(regex, url) is not None
    
    def assign_attributes(self, text, run, attr: dict) -> None:
        font = run.font
        for attribute in attr:
            match str(attribute).capitalize():
                case "Bold":
                    font.bold = True
                case "Size":
                    font.size = Pt(attr['Size'])
                case "Font":
                    font.name = attr['Font']
                case "Italic":
                    font.italic = True
                case "Underline":
                    font.underline = True
                case "Alignment":
                    match attr[attribute]:
                        case "Center":
                            text.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        case "Left":
                            text.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        case "Right":
                            text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                case _: # Ingen- eller ukendt attribute
                    break

    def load_front_page(self) -> None:
        self.heading = self.document.add_heading("", 0)
        self.heading.add_run(self.title, style='Front_page')
        self.heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.heading.bold = True

        # Forklar dette i dokumentationen
        # https://stackoverflow.com/questions/60921603/how-do-i-change-heading-font-face-and-size-in-python-docx
        title_style = self.heading.style
        rFonts = title_style.element.rPr.rFonts
        rFonts.set(qn("w:asciiTheme"), "Times New Roman")

        byline = self.document.add_paragraph("")
        byline.add_run(f"{self.settings['Elev']}\n{self.settings['Klasse']}\n{self.fag.name.capitalize()}", style='Front_page').bold = True
        byline.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if (image := self.config['front_picture']):
            if self.is_url(url=image):
                # TODO: Download image - Slå try/except til igen ved prod.
                #try:
                    from wget import download as wdownload
                    #if not image[-3:] in ["png", "jpg"] and not image[-4:] == "jpeg":
                    #    image = f'{image}.png'
                    self.document.add_picture(wdownload(image, out=f"{__file__[:-8]}/images/"),width=Inches(5.9), height=Inches(2.36)) 
                #except: # Programmet må godt fortsætte
                    #print("Kunne ikke indlæse billede, fortsætter uden...")
            elif os.path.exists(image):
                self.document.add_picture(image, width=Inches(5.9), height=Inches(2.36))
            else:
                print("Kunne ikke indlæse billedet")
                

        self.document.add_page_break()

    def add_style(self, name: str, font: str, size: int) -> None:
        obj_styles = self.document.styles
        obj_font = obj_styles.add_style(name, WD_STYLE_TYPE.CHARACTER).font

        obj_font.name = font
        obj_font.size = Pt(size)

    def load_styles(self) -> None:
        # Font indstillinger
        obj_styles = self.document.styles

        # Forside
        self.add_style(name='Front_page', font=self.data['General']['Font'], size=22)

        # Overskrifter
        self.add_style(name='Overskrift', font=self.data["General"]["Font"], size=self.data["General"]["Headings"]["Size"])

        # Afsnit
        self.add_style(name='Afsnit', font=self.data["General"]["Font"], size=self.data["General"]["Paragraphs"]["Size"])

    def load_header(self) -> None:
        section = self.document.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = f"{self.settings['Elev']} \
            \t{self.fag.name.capitalize()} \
            \t{datetime.today().strftime('%d/%m/%Y')} \
            \n{self.settings['Skole']} \
            \t{self.title} \
            \t{self.settings['Klasse']}"
        
        paragraph.style = self.document.styles["Header"]

    def create(self) -> bool:
        self.document = Document()
        self.data = self.get_content(path=f'{self.fag.value["skabelon"]}')

        # Styles
        self.load_styles()

        # Header
        self.load_header()
        
        # Forside-Tjek
        if self.forside:
            self.load_front_page()

        self.content = self.data['Content']
        for i in self.content:
            match self.content[i]['Type']:
                case "Heading":
                    if not (level := self.content[i].get('Level')):
                        level = 1

                    tmp_heading = self.document.add_heading("", level)
                    run = tmp_heading.add_run(f"{i}", style='Overskrift')

                    if (attributes := self.content[i].get('Attributes')):
                        self.assign_attributes(tmp_heading, run, attributes)

                case "Paragraph":    
                    tmp_paragraph = self.document.add_paragraph("")
                    run = tmp_paragraph.add_run(f"{i}", style='Afsnit')

                    if (attributes := self.content[i].get('Attributes')):
                        self.assign_attributes(tmp_paragraph, run, attributes)
        
        if self.save_path == ".": # Gem lokalt / Bruges til debugging
            self.document.save(f'{self.fag.value["output"]}/{self.title}.docx')
            self.save_path = f'{self.fag.value["output"]}'
        else:
            try:
                print(f'/{self.save_path}.docx')
                self.document.save(f'{self.save_path}/{self.title}.docx')
            except NotADirectoryError:
                self.document.save(f'{self.fag.value["output"]}/{self.title}.docx')
                print("Kunne ikke gemme filen ved den valgte lokation. Gemmer ved ny lokation:")
                print(f'{self.fag.value["output"]}/{self.title}.docx')

        return True

def load_parser() -> dict:
    parser = argparse.ArgumentParser(description="Program til oprettelse af journaler udfra skabeloner i .JSON filer",
                                    formatter_class=argparse.ArgumentDefaultsHelpFormatter)
    parser.add_argument("-t", "--title", help="titel på dokumentet", default="Indsæt Titel Her", required=False)
    parser.add_argument("-f", "--front", action="store_false", help="forside i dokumentet", default=True, required=False)
    parser.add_argument("-p", "--front-picture", help="URL eller sti til lokalt eller online billede til forsiden", required=False)
    parser.add_argument("-o", "--out", help="sti til hvor filen skal gemmes", required=False)
    parser.add_argument("-s", "--subject", help="faget der skal oprettes journal til", required=True)
    args = parser.parse_args()
    return vars(args)

if __name__ == "__main__":
    config = load_parser() # Indlæs argumenter fra CLI
    print(config)
    journal = Journal(config)
    
    if journal.create():
        print("Sucessfully created journal")
        print(f"Opening {(name := journal.title)}")
        
        if platform == "win32":
            os.startfile(f"{journal.save_path}{name}.docx".replace("/", "\\")) 
        elif platform == "darwin":
            # Understøttelse til MacOS
            os.system(f"open -a '/Applications/Microsoft Word.app' '{journal.save_path}/{name}.docx'")
    else:
        raise RuntimeError("Failed to save journal")
    

